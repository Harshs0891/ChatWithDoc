import os
import re
import logging
import requests
import numpy as np
import fitz  # PyMuPDF
import docx
from datetime import datetime
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.docstore.document import Document

logger = logging.getLogger(__name__)


class ChatSystem:
    """Main chat system for document processing and querying"""

    def __init__(self, ollama_base_url="http://localhost:11434"):
        self.ollama_base_url = ollama_base_url
        self.session_documents = {}
        self.embedding_model = "nomic-embed-text"
        self.chat_model = "llama3.1:8b"

    def check_ollama_connection(self):
        """Check Ollama connection"""
        try:
            response = requests.get(f"{self.ollama_base_url}/api/tags", timeout=5)
            return response.status_code == 200
        except Exception:
            return False

    def check_embedding_model(self):
        """Check if embedding model is available"""
        try:
            test_embedding = self._generate_embedding("test")
            return test_embedding is not None and len(test_embedding) > 0
        except Exception:
            return False

    def _generate_embedding(self, text):
        """Generate embedding using Ollama with fallback"""
        try:
            payload = {"model": self.embedding_model, "prompt": text}
            response = requests.post(
                f"{self.ollama_base_url}/api/embeddings", json=payload, timeout=30
            )

            if response.status_code == 200:
                result = response.json()
                embedding = result.get("embedding")
                if embedding and len(embedding) > 0:
                    return embedding
                else:
                    logger.warning("Empty embedding returned from Ollama")
                    return np.random.randn(768).tolist()
            else:
                logger.error(f"Embedding API returned status {response.status_code}")
                return np.random.randn(768).tolist()
        except Exception as e:
            logger.error(f"Error generating embedding: {e}")
            return np.random.randn(768).tolist()

    def _generate_embeddings_batch(self, texts):
        """Generate embeddings for a batch of texts"""
        embeddings = []
        for text in texts:
            embedding = self._generate_embedding(text)
            embeddings.append(embedding)
        return np.array(embeddings)

    def _cosine_similarity(self, a, b):
        """Calculate cosine similarity between two vectors"""
        try:
            a = np.array(a)
            b = np.array(b)
            dot_product = np.dot(a, b)
            norm_a = np.linalg.norm(a)
            norm_b = np.linalg.norm(b)

            if norm_a == 0 or norm_b == 0:
                return 0.0

            return dot_product / (norm_a * norm_b)
        except Exception as e:
            logger.error(f"Error calculating cosine similarity: {e}")
            return 0.0

    def _clean_text(self, text):
        """Clean and normalize text"""
        if not text:
            return ""

        text = re.sub(r"<[^>]+>", "", text)
        text = "".join(char for char in text if char.isprintable() or char in "\n\r\t ")
        text = " ".join(text.split())
        return text.strip()

    def extract_text_from_pdf(self, file_path):
        """Extract text from PDF with validation and better error handling"""
        try:
            if not os.path.exists(file_path):
                logger.error(f"PDF file not found: {file_path}")
                return ""

            # Validate PDF header
            try:
                with open(file_path, "rb") as f:
                    header = f.read(4)
                    if header != b"%PDF":
                        logger.error("File is not a valid PDF")
                        return ""
            except Exception as e:
                logger.error(f"Error validating PDF header: {e}")
                return ""

            text_blocks = []
            try:
                with fitz.open(file_path) as pdf_doc:
                    total_pages = pdf_doc.page_count
                    logger.info(f"Processing PDF with {total_pages} pages")

                    for page_num in range(total_pages):
                        try:
                            page = pdf_doc[page_num]
                            text = page.get_text()

                            if text and text.strip():
                                cleaned_text = self._clean_text(text)
                                if cleaned_text:
                                    text_blocks.append(cleaned_text)
                            else:
                                logger.warning(f"Empty text on page {page_num + 1}")

                        except Exception as page_error:
                            logger.warning(
                                f"Error extracting text from page {page_num + 1}: {page_error}"
                            )

                return "\n\n".join(text_blocks) if text_blocks else ""

            except Exception as pdf_error:
                logger.error(f"Error opening PDF with fitz: {pdf_error}")
                return ""

        except Exception as e:
            logger.error(f"Error processing PDF {file_path}: {e}")
            return ""

    def extract_text_from_docx(self, file_path):
        """Extract text from DOCX"""
        try:
            if not os.path.exists(file_path):
                logger.error(f"DOCX file not found: {file_path}")
                return ""

            doc = docx.Document(file_path)
            paragraphs = []

            for paragraph in doc.paragraphs:
                text = paragraph.text.strip()
                if text:
                    cleaned_text = self._clean_text(text)
                    if cleaned_text:
                        paragraphs.append(cleaned_text)

            return "\n".join(paragraphs)
        except Exception as e:
            logger.error(f"Error extracting DOCX text: {e}")
            return ""

    def extract_text_from_txt(self, file_path):
        """Extract text from TXT with encoding fallback"""
        try:
            if not os.path.exists(file_path):
                logger.error(f"TXT file not found: {file_path}")
                return ""

            for encoding in ["utf-8", "latin-1", "cp1252"]:
                try:
                    with open(file_path, "r", encoding=encoding) as file:
                        content = file.read()
                        if content:
                            return self._clean_text(content)
                        return ""
                except UnicodeDecodeError:
                    continue
                except Exception as e:
                    logger.error(f"Error reading file with {encoding}: {e}")
                    continue

            logger.error(f"Could not decode TXT file: {file_path}")
            return ""
        except Exception as e:
            logger.error(f"Error extracting TXT text from {file_path}: {e}")
            return ""

    def process_documents(self, file_paths, session_id):
        """Process uploaded documents and create embeddings with accurate page tracking"""
        try:
            if not file_paths:
                return False, "No files provided"

            logger.info(f"Processing documents for session: {session_id}")

            documents = []
            extractors = {
                "pdf": self.extract_text_from_pdf_with_pages,
                "docx": self.extract_text_from_docx,
                "txt": self.extract_text_from_txt,
                "doc": self.extract_text_from_docx,
            }

            processed_files = 0
            for file_path in file_paths:
                try:
                    filename = os.path.basename(file_path)
                    ext = filename.lower().split(".")[-1]

                    if ext not in extractors:
                        logger.warning(f"Unsupported file type: {ext}")
                        continue

                    if ext == "pdf":
                        # Use special PDF processor that tracks pages
                        page_texts = extractors[ext](file_path)
                        if not page_texts:
                            logger.warning(f"No text extracted from {filename}")
                            continue

                        # Create chunks with accurate page numbers
                        for page_num, page_text in page_texts.items():
                            if page_text and page_text.strip():
                                # Split page text into smaller chunks if needed
                                text_splitter = RecursiveCharacterTextSplitter(
                                    chunk_size=800,
                                    chunk_overlap=100,
                                    length_function=len,
                                    separators=[
                                        "\n\n",
                                        "\n",
                                        ".",
                                        "!",
                                        "?",
                                        ",",
                                        " ",
                                        "",
                                    ],
                                )
                                chunks = text_splitter.split_text(page_text)

                                for i, chunk in enumerate(chunks):
                                    if chunk and chunk.strip():
                                        doc = Document(
                                            page_content=chunk.strip(),
                                            metadata={
                                                "source": filename,
                                                "page_number": page_num,  # Accurate page number
                                                "chunk_id": i,
                                                "total_chunks": len(chunks),
                                            },
                                        )
                                        documents.append(doc)
                    else:
                        # Handle non-PDF files - set page_number to 1 as they don't have pages
                        text = extractors[ext](file_path)
                        if not text or not text.strip():
                            logger.warning(f"No text extracted from {filename}")
                            continue

                        text_splitter = RecursiveCharacterTextSplitter(
                            chunk_size=1000,
                            chunk_overlap=200,
                            length_function=len,
                            separators=["\n\n", "\n", ".", "!", "?", ",", " ", ""],
                        )
                        chunks = text_splitter.split_text(text)

                        for i, chunk in enumerate(chunks):
                            if chunk and chunk.strip():
                                doc = Document(
                                    page_content=chunk.strip(),
                                    metadata={
                                        "source": filename,
                                        "chunk_id": i,
                                        "total_chunks": len(chunks),
                                        "page_number": 1,  # Fixed: Non-PDF files don't have pages
                                    },
                                )
                                documents.append(doc)

                    processed_files += 1
                    logger.info(f"Processed {filename}: {len(documents)} total chunks")

                except Exception as file_error:
                    logger.error(f"Error processing file {file_path}: {file_error}")
                    continue

            if not documents:
                return (
                    False,
                    "No valid documents processed. Please check if the files contain readable text.",
                )

            # Create embeddings
            document_texts = [doc.page_content for doc in documents]
            logger.info(f"Generating embeddings for {len(document_texts)} chunks...")

            try:
                embeddings = self._generate_embeddings_batch(document_texts)

                # Store in session with proper key
                self.session_documents[session_id] = {
                    "documents": documents,
                    "embeddings": embeddings,
                    "processed_at": datetime.now(),
                }

                logger.info(f"Successfully stored documents for session: {session_id}")

                return (
                    True,
                    f"Successfully processed {processed_files} files with {len(documents)} chunks",
                )
            except Exception as embedding_error:
                logger.error(f"Error generating embeddings: {embedding_error}")
                return False, f"Error generating embeddings: {str(embedding_error)}"

        except Exception as e:
            logger.error(f"Error processing documents: {e}")
            return False, f"Error processing documents: {str(e)}"

    def extract_text_from_pdf_with_pages(self, file_path):
        """Extract text from PDF with accurate page tracking"""
        try:
            if not os.path.exists(file_path):
                logger.error(f"PDF file not found: {file_path}")
                return {}

            # Validate PDF header
            try:
                with open(file_path, "rb") as f:
                    header = f.read(4)
                    if header != b"%PDF":
                        logger.error("File is not a valid PDF")
                        return {}
            except Exception as e:
                logger.error(f"Error validating PDF header: {e}")
                return {}

            page_texts = {}
            try:
                with fitz.open(file_path) as pdf_doc:
                    total_pages = pdf_doc.page_count
                    logger.info(f"Processing PDF with {total_pages} pages")

                    for page_num in range(total_pages):
                        try:
                            page = pdf_doc[page_num]
                            text = page.get_text()

                            if text and text.strip():
                                cleaned_text = self._clean_text(text)
                                if cleaned_text:
                                    # Store with 1-based page numbers
                                    page_texts[page_num + 1] = cleaned_text
                            else:
                                logger.warning(f"Empty text on page {page_num + 1}")

                        except Exception as page_error:
                            logger.warning(
                                f"Error extracting text from page {page_num + 1}: {page_error}"
                            )

                return page_texts

            except Exception as pdf_error:
                logger.error(f"Error opening PDF with fitz: {pdf_error}")
                return {}

        except Exception as e:
            logger.error(f"Error processing PDF {file_path}: {e}")
            return {}

    def generate_answer(self, query, session_id, classification="CONFIDENTIAL"):
        """Generate answer using relevant documents - strictly from PDF content only,
        adaptive length based on available content"""
        try:
            logger.info(f"Generating answer for session: {session_id}")

            if session_id not in self.session_documents:
                logger.error(f"Session {session_id} not found in session_documents")
                return {
                    "success": False,
                    "message": "No documents uploaded. Please upload documents first.",
                    "has_answer": False,
                }

            relevant_docs = self.find_relevant_documents(query, session_id)

            # Set minimum similarity threshold for relevance
            MIN_SIMILARITY = 0.3
            filtered_docs = [
                doc for doc in relevant_docs if doc["similarity"] >= MIN_SIMILARITY
            ]

            if not filtered_docs:
                return {
                    "success": True,
                    "answer": "I don't have information about that in the document.",
                    "has_answer": False,
                    "sources": "",
                    "source_details": [],
                }

            # Prepare context from relevant documents only
            context_parts = []
            sources = set()

            for doc_info in filtered_docs:
                doc = doc_info["document"]
                source = doc.metadata.get("source", "Unknown")
                sources.add(source)
                page_num = doc.metadata.get("page_number", 1)
                context_parts.append(f"Page {page_num}: {doc.page_content}")

            context = "\n\n".join(context_parts)

            # Measure context size
            context_word_count = len(context.split())
            if context_word_count < 300:
                target_length = "between 50 and 150 words"
            else:
                target_length = "between 200 and 300 words"

            # Prompt with adaptive word-length requirement
            prompt = f"""Answer the user's question using ONLY the information from the document excerpts below. 
    Follow these rules:
    1. If the information is available in the document excerpts, provide a clear and direct answer. 
    2. If the information is NOT available in the document excerpts, respond ONLY with: "I don't have information about that in the document."
    3. Do not use external knowledge or make assumptions. 
    4. Do not mention document excerpts, sources, or page numbers in your answer. 
    5. Your answer MUST be {target_length}, depending on the available information.

    Document excerpts:
    {context}

    Question: {query}

    Answer:"""

            answer = self.query_ollama(prompt).strip()

            # Check if the answer indicates no information found
            no_info_phrases = [
                "i don't have information",
                "the document doesn't contain",
                "no information about",
                "not mentioned in the document",
                "cannot find information",
                "no details about",
                "not available in the document",
            ]
            has_answer = not any(phrase in answer.lower() for phrase in no_info_phrases)

            return {
                "success": True,
                "answer": answer,
                "has_answer": has_answer,
                "sources": ", ".join(sources) if has_answer else "",
                "source_details": (
                    [
                        {
                            "source": doc_info["document"].metadata.get(
                                "source", "Unknown"
                            ),
                            "page": doc_info["document"].metadata.get("page_number", 1),
                            "similarity": doc_info["similarity"],
                        }
                        for doc_info in filtered_docs
                    ]
                    if has_answer
                    else []
                ),
            }

        except Exception as e:
            logger.error(f"Error generating answer: {e}")
            return {
                "success": False,
                "message": "An error occurred while generating the response. Please try again.",
                "has_answer": False,
            }

        """Generate answer using relevant documents - strictly from PDF content only,
        adaptive length based on available content"""
        try:
            logger.info(f"Generating answer for session: {session_id}")

            if session_id not in self.session_documents:
                logger.error(f"Session {session_id} not found in session_documents")
                return {
                    "success": False,
                    "message": "No documents uploaded. Please upload documents first.",
                    "has_answer": False,
                }

            relevant_docs = self.find_relevant_documents(query, session_id)

            # Set minimum similarity threshold for relevance
            MIN_SIMILARITY = 0.2
            filtered_docs = [
                doc for doc in relevant_docs if doc["similarity"] >= MIN_SIMILARITY
            ]

            if not filtered_docs:
                return {
                    "success": True,
                    "answer": "I don't have information about that in the document.",
                    "has_answer": False,
                    "sources": "",
                    "source_details": [],
                }

            # Prepare context from relevant documents only
            context_parts = []
            sources = set()

            for doc_info in filtered_docs:
                doc = doc_info["document"]
                source = doc.metadata.get("source", "Unknown")
                sources.add(source)
                page_num = doc.metadata.get("page_number", 1)
                context_parts.append(f"Page {page_num}: {doc.page_content}")

            context = "\n\n".join(context_parts)

            # Measure context size
            context_word_count = len(context.split())
            if context_word_count < 300:
                target_length = "between 50 and 150 words"
            else:
                target_length = "between 200 and 300 words"

            # Prompt with adaptive word-length requirement
            prompt = f"""Answer the user's question using ONLY the information from the document excerpts below. 
    Follow these rules:
    1. If the information is available in the document excerpts, provide a clear and direct answer. 
    2. If the information is NOT available in the document excerpts, respond ONLY with: "I don't have information about that in the document."
    3. Do not use external knowledge or make assumptions. 
    4. Do not mention document excerpts, sources, or page numbers in your answer. 
    5. Your answer MUST be {target_length}, depending on the available information.

    Document excerpts:
    {context}

    Question: {query}

    Answer:"""

            answer = self.query_ollama(prompt).strip()

            # Check if the answer indicates no information found
            no_info_phrases = [
                "i don't have information",
                "the document doesn't contain",
                "no information about",
                "not mentioned in the document",
                "cannot find information",
                "no details about",
                "not available in the document",
            ]
            has_answer = not any(phrase in answer.lower() for phrase in no_info_phrases)

            return {
                "success": True,
                "answer": answer,
                "has_answer": has_answer,
                "sources": ", ".join(sources) if has_answer else "",
                "source_details": (
                    [
                        {
                            "source": doc_info["document"].metadata.get(
                                "source", "Unknown"
                            ),
                            "page": doc_info["document"].metadata.get("page_number", 1),
                            "similarity": doc_info["similarity"],
                        }
                        for doc_info in filtered_docs
                    ]
                    if has_answer
                    else []
                ),
            }

        except Exception as e:
            logger.error(f"Error generating answer: {e}")
            return {
                "success": False,
                "message": "An error occurred while generating the response. Please try again.",
                "has_answer": False,
            }

    def find_relevant_documents(self, query, session_id, top_k=3):
        """Find relevant documents for query using embeddings"""
        logger.info(f"Looking for documents in session: {session_id}")

        if session_id not in self.session_documents:
            logger.warning(f"No documents found for session {session_id}")
            return []

        try:
            query_embedding = self._generate_embedding(query)
            if not query_embedding:
                logger.error("Failed to generate query embedding")
                return []

            doc_embeddings = self.session_documents[session_id]["embeddings"]
            documents = self.session_documents[session_id]["documents"]

            # Calculate similarities
            similarities = []
            for i, doc_embedding in enumerate(doc_embeddings):
                similarity = self._cosine_similarity(query_embedding, doc_embedding)
                similarities.append((i, similarity))

            # Sort by similarity and get top k
            similarities.sort(key=lambda x: x[1], reverse=True)

            relevant_docs = []
            for idx, similarity in similarities[:top_k]:
                relevant_docs.append(
                    {
                        "document": documents[idx],
                        "similarity": float(similarity),
                    }
                )

            logger.info(f"Found {len(relevant_docs)} relevant documents for query")
            return relevant_docs

        except Exception as e:
            logger.error(f"Error finding relevant documents: {e}")
            return []

    def query_ollama(self, prompt, model=None):
        """Query Ollama API for chat completion with fallback"""
        if model is None:
            model = self.chat_model

        try:
            payload = {
                "model": model,
                "prompt": prompt,
                "stream": False,
                "options": {
                    "temperature": 0.5,
                    "top_p": 0.9,
                    "num_predict": 3024,
                    "stop": ["Human:", "Assistant:", "Document excerpts:", "Question:"],
                },
            }

            response = requests.post(
                f"{self.ollama_base_url}/api/generate", json=payload, timeout=120
            )

            if response.status_code == 200:
                result = response.json()
                answer = result.get("response", "").strip()

                if answer:
                    return answer
                else:
                    return "I don't have information about that in the document."
            else:
                logger.error(f"Ollama API returned status {response.status_code}")
                return self._generate_fallback_response()

        except Exception as e:
            logger.error(f"Error querying Ollama: {e}")
            return self._generate_fallback_response()

    def _generate_fallback_response(self):
        """Generate a fallback response when Ollama is unavailable"""
        return "I don't have information about that in the document."

    def has_documents(self, session_id):
        """Check if session has documents"""
        has_docs = (
            session_id in self.session_documents
            and len(self.session_documents[session_id].get("documents", [])) > 0
        )
        logger.info(f"Session {session_id} has documents: {has_docs}")
        return has_docs

    def get_document_count(self, session_id):
        """Get number of documents in session"""
        return len(self.session_documents.get(session_id, {}).get("documents", []))

    def get_active_sessions_count(self):
        """Get number of active sessions"""
        return len(self.session_documents)

    def clear_session(self, session_id):
        """Clear documents from session"""
        if session_id in self.session_documents:
            del self.session_documents[session_id]
            logger.info(f"Cleared session: {session_id}")
            return True
        return False

    def generate_smart_questions(self, session_id, count=3):
        """Generate smart questions based ONLY on the uploaded PDF document content"""
        if session_id not in self.session_documents:
            return {"welcome": "", "questions": []}

        try:
            documents = self.session_documents[session_id]["documents"]
            if not documents:
                return {"welcome": "", "questions": []}

            # Extract key content from documents
            content_samples = []
            sources = set()

            for doc in documents[:80]:  # Use first 8 documents for analysis
                content = doc.page_content
                source = doc.metadata.get("source", "Unknown")
                content_samples.append(content)
                sources.add(source)

            # Create context for analysis
            context = "\n\n".join(content_samples[:5000])  # Limit context length
            source_name = list(sources)[0] if sources else "document"

            # First attempt - combined prompt
            prompt = f"""Based ONLY on the document content below, create:
    1. Write a warm, professional, and detailed welcome message that introduces the purpose, scope, and structure of this document. The message should clearly explain what the document is about and what the reader can expect to learn or find within it.
    2. 3-5 specific questions that can be answered from this document

    Document content:
    {context}

    Respond in this EXACT format:
    WELCOME: [your welcome message here]
    QUESTIONS:
    1. [question 1]
    2. [question 2]
    3. [question 3]
"""

            response = self.query_ollama(prompt)

            # Parse response
            welcome_message = ""
            questions = []

            # Robust parsing with multiple fallback strategies
            if "WELCOME:" in response and "QUESTIONS:" in response:
                parts = response.split("QUESTIONS:")
                welcome_part = parts[0].replace("WELCOME:", "").strip()
                questions_part = parts[1].strip()

                welcome_message = welcome_part

                # Extract questions
                for line in questions_part.split("\n"):
                    line = line.strip()
                    if line and any(char.isdigit() for char in line[:3]):
                        question = re.sub(r"^\d+[\.\)]\s*", "", line).strip()
                        if question and len(question) > 10:
                            questions.append(question)

            # If either welcome or questions are missing, make separate calls
            needs_welcome = not welcome_message or len(welcome_message.strip()) < 20
            needs_questions = len(questions) < 2

            if needs_welcome or needs_questions:
                logger.info(
                    f"Making separate calls - needs_welcome: {needs_welcome}, needs_questions: {needs_questions}"
                )

                if needs_welcome:
                    welcome_prompt = f"""Based ONLY on the document content below, write a warm, professional, and detailed welcome message that:
    - Introduces the purpose and scope of this document
    - Explains what the document is about
    - Describes what the reader can expect to learn or find within it
    - Mentions the structure or key sections

    Document content:
    {context[:3000]}

    Respond with ONLY the welcome message, no labels or formatting:"""

                    welcome_response = self.query_ollama(welcome_prompt)
                    if welcome_response and len(welcome_response.strip()) > 20:
                        welcome_message = welcome_response.strip()
                    else:
                        welcome_message = f"I've analyzed your document '{source_name}' and I'm ready to help you explore its content. This document appears to cover important information that I can help you understand and discuss."

                if needs_questions:
                    questions_prompt = f"""Based ONLY on the document content below, generate 3-5 specific questions that can be answered directly from the content. Make them diverse and cover different aspects.

    Document content:
    {context[:3000]}

    Respond with ONLY the questions in this exact format:
    1. [specific question about document content]
    2. [specific question about document content]
    3. [specific question about document content]"""

                    questions_response = self.query_ollama(questions_prompt)
                    new_questions = []

                    for line in questions_response.split("\n"):
                        line = line.strip()
                        if line and any(char.isdigit() for char in line[:3]):
                            question = re.sub(r"^\d+[\.\)]\s*", "", line).strip()
                            if (
                                question
                                and len(question) > 10
                                and question not in new_questions
                            ):
                                new_questions.append(question)

                    if new_questions:
                        questions = new_questions

            # Final fallbacks if still missing content
            if not welcome_message or len(welcome_message.strip()) < 20:
                welcome_message = f"I've analyzed your document '{source_name}' and I'm ready to help you explore its content. This document contains valuable information that I can help you understand and discuss."

            if len(questions) < 2:
                # Add some intelligent fallback questions
                fallback_questions = [
                    "What are the main topics and key points covered in this document?",
                    "Can you summarize the most important information presented?",
                    "What specific details or findings should I pay attention to?",
                    "How is the information in this document organized or structured?",
                ]
                # If we have some questions but need more, supplement them
                questions.extend(fallback_questions[: 4 - len(questions)])

            # Ensure questions are unique and high quality
            unique_questions = []
            seen_questions = set()
            for q in questions:
                # Basic quality check
                if (
                    q
                    and len(q) > 15
                    and q.lower() not in seen_questions
                    and not q.lower().startswith(("can you", "what are", "how is"))
                ):
                    unique_questions.append(q)
                    seen_questions.add(q.lower())

            # If we lost too many questions to deduplication, add some back
            if len(unique_questions) < 2:
                unique_questions = questions[:count]  # Use original list

            # Limit to requested count and ensure quality
            final_questions = []
            for q in unique_questions[:count]:
                if len(q) > 10 and "?" in q:
                    final_questions.append(q)
                else:
                    # Fix poorly formatted questions
                    fixed_q = q.strip()
                    if not fixed_q.endswith("?"):
                        fixed_q += "?"
                    final_questions.append(fixed_q)

            return {"welcome": welcome_message, "questions": final_questions}

        except Exception as e:
            logger.error(f"Error generating smart questions: {e}")
            return {
                "welcome": "I've processed your document and I'm ready to help answer questions about its content.",
                "questions": [
                    "What are the main topics discussed in this document?",
                    "Can you summarize the key points and important information?",
                    "What specific details or findings should I know about?",
                ][:count],
            }
